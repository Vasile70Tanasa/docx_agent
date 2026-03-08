"""filler.py – DOCX Form Filler with Visual Width Preservation

When a placeholder (_____ or .....) is replaced by a value, the
replacement is fitted to the *exact same visual width* using PIL font
metrics.  This prevents Word from reflowing lines and shifting pagination.

API (unchanged from v3):
    fill_document(docx_in, docx_out, fields, tables, mapping, data) → stats
"""
from __future__ import annotations

import json, os, re
from difflib import SequenceMatcher
from functools import lru_cache
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from PIL import ImageFont


# ── 1. FONT WIDTH ENGINE ─────────────────────────────────────────────────────

_FONT_DIRS = [
    os.path.join(os.environ.get('WINDIR', r'C:\Windows'), 'Fonts'),
    '/usr/share/fonts/truetype',
]

_FONT_MAP: Dict[str, str] = {}          # populated lazily


def _find_font_file(font_name: str) -> str:
    """Resolve a font family name to a .ttf path (Windows-centric)."""
    key = font_name.lower().replace(' ', '')
    if key in _FONT_MAP:
        return _FONT_MAP[key]

    for d in _FONT_DIRS:
        if not os.path.isdir(d):
            continue
        for fn in os.listdir(d):
            if fn.lower().endswith(('.ttf', '.ttc')):
                stem = fn.rsplit('.', 1)[0].lower().replace(' ', '')
                _FONT_MAP[stem] = os.path.join(d, fn)

    # Try exact, then substring
    for stem, path in _FONT_MAP.items():
        if stem == key:
            return path
    for stem, path in _FONT_MAP.items():
        if key in stem or stem in key:
            return path
    return ''


@lru_cache(maxsize=32)
def _load_font(font_name: str, size_pt: float) -> ImageFont.FreeTypeFont | None:
    path = _find_font_file(font_name)
    if not path:
        return None
    try:
        return ImageFont.truetype(path, int(size_pt))
    except Exception:
        return None


def _measure(text: str, font_name: str, size_pt: float) -> float:
    """Visual width (pixels) of *text* in the given font."""
    font = _load_font(font_name, size_pt)
    if font is None:
        return float(len(text))          # graceful fallback: 1 px per char
    return font.getlength(text)


def _run_font(run) -> Tuple[str, float]:
    """Extract (font_name, size_pt) from a python-docx Run.

    Walks the inheritance chain: run -> paragraph style -> document defaults
    so that pages with 11pt inherit 11pt, not a hardcoded 12pt fallback.
    """
    _WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    name = run.font.name
    size = run.font.size  # EMU or None

    # Walk up to paragraph style
    if not name or not size:
        try:
            para_elem = run._r.getparent()
            doc = run._parent._parent
            pPr = para_elem.find(f'{{{_WML_NS}}}pPr')
            pstyle_id = None
            if pPr is not None:
                pStyle = pPr.find(f'{{{_WML_NS}}}pStyle')
                if pStyle is not None:
                    pstyle_id = pStyle.get(f'{{{_WML_NS}}}val')
            if pstyle_id and pstyle_id in doc.styles:
                st = doc.styles[pstyle_id]
                if not name and st.font.name:
                    name = st.font.name
                if not size and st.font.size:
                    size = st.font.size
        except Exception:
            pass

    # Fall back to Normal / document default
    if not name or not size:
        try:
            normal = run._parent._parent.styles['Normal']
            if not name and normal.font.name:
                name = normal.font.name
            if not size and normal.font.size:
                size = normal.font.size
        except Exception:
            pass

    pt = size / 12700 if size else 11.0   # 11pt safer default than 12
    return name or 'Times New Roman', pt


def _fit_to_field(replacement: str, placeholder: str,
                  para_text: str, placeholder_end: int,
                  size_pt: float, font_name: str = 'Times New Roman') -> Tuple[str, float]:
    """Fit *replacement* into the placeholder field.

    - Normalise whitespace with regular spaces (allows Word to word-wrap naturally).
    - UNDERSCORE fields: pad with '_' to same character count (underscores are wide).
    - DOT fields: pad with '.' to match VISUAL width (dots are narrow; character padding
      adds too many dots, pushing text like 'ofertant' to the next line).
    - If replacement is LONGER: insert as-is.
    - No font changes.
    """
    had_leading = replacement.startswith(' ') or replacement.startswith('\t')
    replacement = ' '.join(replacement.split())
    if had_leading and replacement:
        replacement = ' ' + replacement

    if '_' in placeholder:
        n_ph  = len(placeholder)
        n_rep = len(replacement)
        if n_rep < n_ph:
            return replacement + '_' * (n_ph - n_rep), size_pt
        return replacement, size_pt

    if '.' in placeholder:
        text_w = _measure(replacement, font_name, size_pt)
        target_w = _measure(placeholder, font_name, size_pt)
        if text_w < target_w:
            pad_w = _measure('.', font_name, size_pt)
            if pad_w > 0:
                n_pads = max(0, round((target_w - text_w) / pad_w))
                return replacement + '.' * n_pads, size_pt
        return replacement, size_pt

    return replacement, size_pt


# Kept for compatibility with _replace_dots
def _fit_visual(replacement: str, placeholder: str,
                font_name: str, size_pt: float) -> str:
    text, _ = _fit_to_field(replacement, placeholder, '', 0, size_pt)
    return text


# ── 2. DOCUMENT TRAVERSAL ────────────────────────────────────────────────────

def _iter_blocks(parent):
    """Yield Paragraph / Table children in document order."""
    elm = parent.element.body if isinstance(parent, DocxDocument) else parent._tc
    for child in elm.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            yield Paragraph(child, parent)
        elif tag == 'tbl':
            yield Table(child, parent)


_WML = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _section_has(section, kind: str) -> bool:
    """True when sectionPr has an explicit header/footerReference."""
    return len(section._sectPr.findall(f'{{{_WML}}}{kind}Reference')) > 0


def _build_lookup(doc: Document) -> Dict[str, Paragraph]:
    """Map parser location strings to live Paragraph objects."""
    lk: Dict[str, Paragraph] = {}
    pi = ti = 0
    for blk in _iter_blocks(doc):
        if isinstance(blk, Paragraph):
            lk[f'body/p:{pi}'] = blk
            pi += 1
        elif isinstance(blk, Table):
            for ri, row in enumerate(blk.rows):
                for ci, cell in enumerate(row.cells):
                    cpi = 0
                    for b2 in _iter_blocks(cell):
                        if isinstance(b2, Paragraph):
                            lk[f'body/t:{ti}/r:{ri}/c:{ci}/p:{cpi}'] = b2
                            cpi += 1
            ti += 1
    for si, sec in enumerate(doc.sections):
        if _section_has(sec, 'header'):
            for pi2, p in enumerate(sec.header.paragraphs):
                lk[f'sec:{si}/header/p:{pi2}'] = p
        if _section_has(sec, 'footer'):
            for pi2, p in enumerate(sec.footer.paragraphs):
                lk[f'sec:{si}/footer/p:{pi2}'] = p
    return lk


# ── 3. SPAN REPLACEMENT ─────────────────────────────────────────────────────

def _replace_span(para: Paragraph, start: int, end: int, replacement: str,
                  font_name: str, size_pt: float,
                  nb_spaces: bool = False) -> bool:
    """Replace characters [start:end] in *para* with *replacement*.
    Pads to visual width; no font changes.
    """
    runs = para.runs
    if not runs:
        return False

    offsets: List[Tuple[int, int]] = []
    pos = 0
    for r in runs:
        t = r.text or ''
        offsets.append((pos, pos + len(t)))
        pos += len(t)

    if start < 0 or end > pos or start >= end:
        return False

    full      = para.text or ''
    orig_text = full[start:end]

    fitted_text, _ = _fit_to_field(
        replacement, orig_text, full, end, size_pt, font_name
    )

    if nb_spaces:
        fitted_text = fitted_text.replace(' ', '\u00a0')

    placed = False
    for i, (rs, re_) in enumerate(offsets):
        if re_ <= start or rs >= end:
            continue
        run = runs[i]
        txt = run.text or ''
        ls  = max(0, start - rs)
        le  = min(len(txt), end - rs)
        if not placed:
            run.text = txt[:ls] + fitted_text + txt[le:]
            placed = True
        else:
            run.text = txt[:ls] + txt[le:]
    return placed


_RE_DOTS = re.compile(r'\.{5,}')


def _replace_dots(para: Paragraph, replacement: str,
                  start: int, end: int,
                  font_name: str, size_pt: float) -> bool:
    """Find the dots pattern nearest to [start,end] and replace it."""
    full = para.text or ''
    hit = None
    for m in _RE_DOTS.finditer(full):
        if m.start() < end and m.end() > start:
            hit = m
            break
    if hit is None:
        for m in _RE_DOTS.finditer(full):
            hit = m
            break
    if hit:
        return _replace_span(para, hit.start(), hit.end(), replacement,
                             font_name, size_pt)
    return False


def _replace_checkbox(para: Paragraph, start: int, end: int) -> bool:
    """Replace |_| → |x| (same visual width, no fitting needed)."""
    runs = para.runs
    if not runs:
        return False
    offsets, pos = [], 0
    for r in runs:
        t = r.text or ''
        offsets.append((pos, pos + len(t)))
        pos += len(t)
    if start < 0 or end > pos or start >= end:
        return False
    # Direct span
    placed = False
    for i, (rs, re_) in enumerate(offsets):
        if re_ <= start or rs >= end:
            continue
        run = runs[i]
        txt = run.text or ''
        ls = max(0, start - rs)
        le = min(len(txt), end - rs)
        if not placed:
            run.text = txt[:ls] + '|x|' + txt[le:]
            placed = True
        else:
            run.text = txt[:ls] + txt[le:]
    if placed:
        return True
    # Fallback: regex
    full = para.text or ''
    if '|_|' in full:
        return _replace_dots(para, '|x|', start, end, 'Arial', 11)
    if '☐' in full:
        for r in runs:
            if '☐' in (r.text or ''):
                r.text = r.text.replace('☐', '☑', 1)
                return True
    return False


# ── 4. VALUE FORMATTING ─────────────────────────────────────────────────────

_RE_ISO = re.compile(r'^(\d{4})-(\d{2})-(\d{2})$')


def _fmt(value: Any) -> str:
    """JSON value → display string."""
    if value is None:
        return ''
    if isinstance(value, bool):
        return 'da' if value else 'nu'
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, dict):
        return ''
    if isinstance(value, list):
        if not value:
            return ''
        if all(isinstance(x, (str, int, float, bool)) for x in value):
            return '; '.join(str(x) for x in value)
        return ''
    s = str(value).strip()
    m = _RE_ISO.match(s)
    if m:
        return f'{m.group(3)}/{m.group(2)}/{m.group(1)}'
    if s.startswith('['):
        try:
            parsed = json.loads(s)
            if isinstance(parsed, list) and all(
                isinstance(x, (str, int, float, bool)) for x in parsed
            ):
                return '; '.join(str(x) for x in parsed)
        except Exception:
            pass
    return str(value)


def _get(data: Dict, key: str) -> Any:
    """Get value, supporting dotted paths."""
    if key in data:
        return data[key]
    cur: Any = data
    for p in key.split('.'):
        if not isinstance(cur, dict) or p not in cur:
            return None
        cur = cur[p]
    return cur


# ── 5. TEXT MATCHING HELPERS ─────────────────────────────────────────────────

def _fold(s: str) -> str:
    for a, b in [('ă','a'),('â','a'),('î','i'),('ș','s'),('ş','s'),('ț','t'),('ţ','t')]:
        s = s.replace(a, b)
    return s


def _norm(s: str) -> str:
    s = _fold((s or '').lower()).replace('_', ' ')
    return re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', ' ', s)).strip()


def _sim(a: str, b: str) -> float:
    return SequenceMatcher(None, _norm(a), _norm(b)).ratio()


def _option_match(label: str, value: str) -> bool:
    nl, nv = _norm(label), _norm(str(value))
    if not nl or not nv:
        return False
    return nl == nv or nv in nl or nl in nv or _sim(label, str(value)) > 0.72


def _last_token(s: str) -> str:
    """Last token from context (after comma or whole string)."""
    s = (s or '').strip()
    if ',' in s:
        return s.split(',')[-1].strip()
    return s


def _adjust_replacement(repl: str, ctx_before: str) -> str:
    """Fix spacing and prefix duplication.

    - Add leading space when context ends with text (no space), e.g. 'Data completării' + '08/10/2025'
    - If value duplicates what template already has: remove the duplicated part, keep only input value.
      E.g. template "dl." + placeholder, value "dl. Popa Mihai" -> output " Popa Mihai" -> "dl. Popa Mihai"
      E.g. template "dna" + placeholder, value "dna. Ionescu Ana" -> output ". Ionescu Ana" -> "dna. Ionescu Ana"
    """
    if not repl:
        return repl
    ctx = (ctx_before or '').strip()
    repl = repl.strip()

    # 1. Remove duplicated prefix: template has T, value has "T X" -> output only "X" with proper spacing
    stripped_prefix = False
    if ctx and repl:
        token = _last_token(ctx)
        if token:
            rest = None
            # repl can start with "token. " or "token " (token may be "dl." or "dna")
            if token.endswith('.'):
                prefix = token + ' '
                if repl.lower().startswith(prefix.lower()):
                    rest = repl[len(prefix):].lstrip()
            else:
                for sep in ('. ', ' '):
                    prefix = token + sep
                    if repl.lower().startswith(prefix.lower()):
                        rest = repl[len(prefix):].lstrip()
                        break
            if rest is not None:
                # Pad: if ctx ends with ".", add " "; else add ". " (e.g. "dna" needs ". " for "dna. X")
                pad = ' ' if ctx.rstrip().endswith('.') else '. '
                repl = pad + rest
                stripped_prefix = True

    # 2. Add leading space when context ends with letter (no space before placeholder)
    if repl and not repl.startswith(' ') and not repl.startswith('.'):
        if stripped_prefix or (ctx and not ctx.endswith(' ')):
            repl = ' ' + repl

    return repl


# ── 6. TABLE FILL ───────────────────────────────────────────────────────────

def _col_score(col_hdr: str, dict_key: str) -> float:
    s = _sim(col_hdr, dict_key)
    nh = _norm(col_hdr)
    toks = [t for t in re.split(r'\W+', _norm(dict_key)) if len(t) >= 4]
    if toks:
        s = max(s, sum(1 for t in toks if t in nh) / len(toks) * 0.75)
    return s


def _is_nr_col(hdr: str) -> bool:
    n = _norm(hdr)
    return n in ('nr', 'nr crt') or ('nr' in n and 'crt' in n and len(hdr.strip()) < 15)


def _write_cell(cell: _Cell, text: str, font_name: str, size_pt: float) -> None:
    """Write text into cell, preserving visual width of existing content."""
    if not cell.paragraphs:
        return
    para = cell.paragraphs[0]
    # Find a placeholder run
    for r in para.runs:
        if '_' in r.text or '.' in r.text:
            fitted = _fit_visual(text, r.text, font_name, size_pt)
            r.text = fitted
            return
    # No placeholder — fit to existing content
    existing = ''.join(r.text or '' for r in para.runs)
    if existing.strip():
        fitted = _fit_visual(text, existing, font_name, size_pt)
    else:
        fitted = text  # truly empty cell — just write
    if para.runs:
        para.runs[-1].text = fitted
    else:
        para.add_run(fitted)


def _fill_table(doc: Document, meta: Dict, value: Any, stats: Dict,
                font_name: str, size_pt: float) -> None:
    """Fill table data rows from a JSON array."""
    tables = doc.tables
    ti = meta['table_index']
    if ti >= len(tables):
        return
    table = tables[ti]
    headers: List[str] = meta['col_headers']
    data_rows: List[int] = meta['data_rows']

    # Parse stringified JSON
    if isinstance(value, str):
        s = value.strip()
        if s.startswith(('[', '{')):
            try: value = json.loads(s)
            except Exception: pass
    if not isinstance(value, list):
        value = [value] if value else []

    for ei, entry in enumerate(value):
        if ei >= len(data_rows):
            break
        ri = data_rows[ei]
        if ri >= len(table.rows):
            break
        row = table.rows[ri]

        if isinstance(entry, str):
            try: entry = json.loads(entry)
            except Exception: pass

        if isinstance(entry, dict):
            ekeys = list(entry.keys())
            for ci, cell in enumerate(row.cells):
                if ci >= len(headers):
                    continue
                if _is_nr_col(headers[ci]):
                    _write_cell(cell, str(ei + 1), font_name, size_pt)
                    print(f"DEBUG: Wrote NR_COL {str(ei + 1)} to cell {ci}")
                    stats['filled'] += 1
                    continue
                best_k, best_s = None, 0.0
                for ek in ekeys:
                    sc = _col_score(headers[ci], ek)
                    if sc > best_s:
                        best_s, best_k = sc, ek
                if best_k and best_s > 0.30:
                    _write_cell(cell, _fmt(entry[best_k]), font_name, size_pt)
                    print(f"DEBUG: Wrote {best_k} to cell {ci}")
                    stats['filled'] += 1
        elif isinstance(entry, str) and row.cells:
            _write_cell(row.cells[0], entry, font_name, size_pt)
            print(f"DEBUG: Wrote fallback string to cell 0")
            stats['filled'] += 1


# ── 7. MAIN FILL ────────────────────────────────────────────────────────────

def _detect_font(para: Paragraph, start: int, end: int) -> Tuple[str, float]:
    """Get font info from the run that contains the placeholder.

    IMPORTANT: We skip runs whose font.size was explicitly set to a small value
    by a previous shrink-to-fit replacement (< 8pt), because that would
    contaminate all subsequent placeholders in the same paragraph.
    Instead we walk up to the paragraph style as authoritative source.
    """
    _WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # 1. Try to get font from the run that contains the placeholder,
    #    but ONLY if the run has no size override (or a plausible size >= 8pt).
    pos = 0
    for r in para.runs:
        t = r.text or ''
        rend = pos + len(t)
        if rend > start and pos < end:
            # Check if this run has an explicit size that is ours (tiny)
            sz = r.font.size
            if sz is None or sz / 12700 >= 8.0:
                # Safe to use this run's font
                return _run_font(r)
            # else: this run was shrunk by us — fall through below
        pos = rend

    # 2. Get font from paragraph style (immune to run-level overrides)
    try:
        doc = para._p.getparent().getparent()   # body -> document
        # Try paragraph style
        pPr = para._p.find(f'{{{_WML_NS}}}pPr')
        if pPr is not None:
            pStyle = pPr.find(f'{{{_WML_NS}}}pStyle')
            if pStyle is not None:
                sid = pStyle.get(f'{{{_WML_NS}}}val')
                from docx import Document as _DocxModule
                # Access styles via the parent document
                from docx.oxml.ns import qn
                styles_elem = para._p.getroottree().getroot().find(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styles')
                if styles_elem is not None:
                    for st_elem in styles_elem:
                        sid_attr = st_elem.get(f'{{{_WML_NS}}}styleId')
                        if sid_attr == sid:
                            rpr = st_elem.find(f'.//{{{_WML_NS}}}rPr')
                            if rpr is not None:
                                sz_elem = rpr.find(f'{{{_WML_NS}}}sz')
                                if sz_elem is not None:
                                    half_pt = int(sz_elem.get(f'{{{_WML_NS}}}val', '0'))
                                    if half_pt > 0:
                                        name_elem = rpr.find(f'{{{_WML_NS}}}rFonts')
                                        fname = 'Times New Roman'
                                        if name_elem is not None:
                                            fname = (name_elem.get(f'{{{_WML_NS}}}ascii')
                                                     or name_elem.get(f'{{{_WML_NS}}}hAnsi')
                                                     or 'Times New Roman')
                                        return fname, half_pt / 2.0
    except Exception:
        pass

    # 3. Last resort: first non-tiny run
    for r in para.runs:
        sz = r.font.size
        if sz is None or sz / 12700 >= 8.0:
            return _run_font(r)

    return ('Times New Roman', 11.0)


def _has_dots(raw_span: str) -> bool:
    return bool(raw_span) and (bool(_RE_DOTS.search(raw_span)) or raw_span.count('.') >= 5)


def fill_document(
    docx_in: str,
    docx_out: str,
    fields: List[Dict],
    tables: List[Dict],
    mapping: Dict[str, Any],
    data: Dict[str, Any],
) -> Dict[str, int]:
    """Fill DOCX form and save.  Returns {filled, skipped, failed}."""
    doc = Document(docx_in)
    lookup = _build_lookup(doc)
    stats = {'filled': 0, 'skipped': 0, 'failed': 0}

    # ── TEXT / DATE fields ───────────────────────────────────────────────────
    # Group by paragraph (multiple fields in same para must be applied right→left)
    pending: Dict[str, List[Tuple[int, int, str, Dict]]] = {}

    array_tracker: Dict[str, int] = {}

    for f in fields:
        if f['field_type'] == 'CHECKBOX':
            continue
        m = mapping.get(f['field_id'], {})
        jk = m.get('json_key')
        if not jk:
            stats['skipped'] += 1
            continue
        val = _get(data, jk)
        if isinstance(val, str) and val.strip().startswith('['):
            try:
                parsed = json.loads(val)
                if isinstance(parsed, list):
                    val = parsed
            except Exception:
                pass
                
        if val is None:
            stats['skipped'] += 1
            continue
        if f['location'] not in lookup:
            stats['failed'] += 1
            continue
            
        if isinstance(val, list) and not m.get('extracted_value'):
            if jk not in array_tracker:
                array_tracker[jk] = 0
            idx = array_tracker[jk]
            if idx < len(val):
                repl = _fmt(val[idx])
                array_tracker[jk] += 1
            else:
                repl = "" # Exhausted
        else:
            repl = _fmt(val)
            if m.get('extracted_value'):
                repl = str(m['extracted_value']).strip()
            
        if not repl:
            stats['skipped'] += 1
            continue

        pending.setdefault(f['location'], []).append(
            (int(f['start']), int(f['end']), repl, f, m)
        )

    # ── Detect "linked" fields: same json_key, consecutive paragraphs,
    #    one ends at para tail and next starts at para head.
    #    Keep only the one with the longest placeholder; skip the shorter one.
    _skip_ids: set = set()
    _by_key: Dict[str, List[Tuple[str, int, int, Dict, Dict]]] = {}
    for loc, reps in pending.items():
        for start, end, repl, fld, m in reps:
            jk = m.get('json_key', '')
            if jk:
                _by_key.setdefault(jk, []).append((loc, start, end, fld, m))
    for jk, entries in _by_key.items():
        if len(entries) < 2:
            continue
        entries.sort(key=lambda e: e[0])  # sort by location
        for i in range(len(entries) - 1):
            loc_a, s_a, e_a, fld_a, _ = entries[i]
            loc_b, s_b, e_b, fld_b, _ = entries[i + 1]
            seq_a = fld_a.get('seq', -1)
            seq_b = fld_b.get('seq', -1)
            ft_a = fld_a.get('full_text', '')
            if seq_b - seq_a != 1:
                continue
            if e_a != len(ft_a) or s_b != 0:
                continue
            # Linked pair — skip the shorter placeholder
            span_a = e_a - s_a
            span_b = e_b - s_b
            skip_fld = fld_a if span_a < span_b else fld_b
            _skip_ids.add(skip_fld['field_id'])

    # Remove skipped linked fields from pending
    for loc in list(pending):
        pending[loc] = [(s, e, r, f, m) for s, e, r, f, m in pending[loc]
                        if f['field_id'] not in _skip_ids]
        if not pending[loc]:
            del pending[loc]

    for loc, reps in pending.items():
        para = lookup.get(loc)
        if not para:
            continue
        reps.sort(key=lambda x: -x[0])     # right-to-left
        for start, end, repl, fld, m in reps:
            # Use actual document text as ctx (more reliable than cached mapping)
            para_text = para.text or ''
            ctx_before = para_text[:start].rstrip()
            if not ctx_before:
                ctx_before = m.get('ctx_before') or fld.get('ctx_before', '')
            repl = _adjust_replacement(repl, ctx_before)
            # Ensure space: if char immediately before placeholder is not whitespace, add leading space
            if repl and start > 0 and len(para_text) >= start:
                char_before = para_text[start - 1]
                if char_before not in ' \t\n' and not repl.startswith(' ') and not repl.startswith('.'):
                    repl = ' ' + repl
            fn, sz = _detect_font(para, start, end)
            raw = fld.get('raw_span', '')
            if _has_dots(raw):
                ok = _replace_dots(para, repl, start, end, fn, sz)
            else:
                ok = _replace_span(para, start, end, repl, fn, sz,
                                   nb_spaces='\t' in para_text[:start])
            stats['filled' if ok else 'failed'] += 1

    # ── CHECKBOX groups ──────────────────────────────────────────────────────
    groups: Dict[str, List[Dict]] = {}
    for f in fields:
        if f['field_type'] == 'CHECKBOX' and f.get('group_id'):
            groups.setdefault(f['group_id'], []).append(f)

    for gid, gfields in groups.items():
        m = mapping.get(gid, {})
        jk = m.get('json_key')
        if not jk:
            for dk, dv in data.items():
                if dv is None:
                    continue
                vs = str(dv).strip()
                for f in gfields:
                    opt = f.get('option_label') or f.get('label', '')
                    if _option_match(opt, vs):
                        jk = dk
                        break
                if jk:
                    break
        if not jk:
            stats['skipped'] += len(gfields)
            continue
        val = _get(data, jk)
        if val is None:
            stats['skipped'] += len(gfields)
            continue
        vs = str(val)
        for f in gfields:
            para = lookup.get(f['location'])
            if not para:
                stats['failed'] += 1
                continue
            opt = f.get('option_label') or f.get('label', '')
            if _option_match(opt, vs):
                ok = _replace_checkbox(para, int(f['start']), int(f['end']))
                stats['filled' if ok else 'failed'] += 1
            else:
                stats['skipped'] += 1

    # ── TABLE arrays ─────────────────────────────────────────────────────────
    # Detect a default font from first body paragraph run
    default_fn, default_sz = 'Times New Roman', 12.0
    for f in fields:
        p = lookup.get(f.get('location', ''))
        if p and p.runs:
            default_fn, default_sz = _run_font(p.runs[0])
            break

    for t in tables:
        m = mapping.get(t['field_id'], {})
        jk = m.get('json_key')
        if not jk:
            stats['skipped'] += 1
            continue
        val = _get(data, jk)
        if val is None:
            stats['skipped'] += 1
            continue
        _fill_table(doc, t, val, stats, default_fn, default_sz)

    doc.save(docx_out)
    return stats


# ── CLI ──────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    import argparse

    ap = argparse.ArgumentParser(description='Fill DOCX from mapping + data')
    ap.add_argument('--docx', default='sample_forms.docx')
    ap.add_argument('--parser', default='parser_enriched_full.json')
    ap.add_argument('--mapping', default='mapping.json')
    ap.add_argument('--data', default='input_date_expanded.json')
    ap.add_argument('--out', default='filled.docx')
    args = ap.parse_args()

    sd = os.path.dirname(os.path.abspath(__file__))
    dp = args.docx if os.path.isabs(args.docx) else os.path.join(sd, args.docx)
    pp = args.parser if os.path.isabs(args.parser) else os.path.join(sd, args.parser)
    mp = args.mapping if os.path.isabs(args.mapping) else os.path.join(sd, args.mapping)
    datp = args.data if os.path.isabs(args.data) else os.path.join(sd, args.data)
    op = args.out if os.path.isabs(args.out) else os.path.join(sd, args.out)

    with open(pp, 'r', encoding='utf-8') as f:
        parsed = json.load(f)
    with open(mp, 'r', encoding='utf-8') as f:
        mapping = json.load(f)
    with open(datp, 'r', encoding='utf-8') as f:
        data_json = json.load(f)

    s = fill_document(dp, op, parsed['fields'], parsed.get('tables', []), mapping, data_json)
    print(f'Filled: {s["filled"]}  Skipped: {s["skipped"]}  Failed: {s["failed"]}')
    print(f'Output: {op}')
