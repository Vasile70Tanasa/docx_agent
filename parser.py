"""parser.py - DOCX Form Field Extractor (v3)

Scans a DOCX document and returns all fillable fields with their labels and context.

Field types detected:
  TEXT     - underscore (___) or dots (....) placeholders
  DATE     - date pattern __/__/__
  CHECKBOX - |_| option boxes (grouped when consecutive)

Label extraction priority:
  1. (hint in parens) in suffix on same line
  2. (hint in parens) anywhere in the paragraph
  3. Next paragraph if it looks like a pure label hint
  4. CHECKBOX: text after |_| on the same line
  5. Last meaningful chunk of the prefix
  6. suffix_start for fields at start of paragraph

Also produces TableMeta objects for tables that have empty data rows,
to be filled from JSON array values.

IMPORTANT: field_id is stable — does NOT include label in hash.
"""
from __future__ import annotations

import re
import hashlib
from dataclasses import dataclass, asdict
from typing import Any, Dict, Iterator, List, Optional, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


# ── Placeholder regexes ──────────────────────────────────────────────────────
RE_DATE      = re.compile(r'_{2,}\s*/\s*_{2,}\s*/\s*_{2,}')
RE_UND       = re.compile(r'_{3,}(?:(?:[ \t]{1,3}|\n\s*|/|\s*/\s*)_{2,})*')
RE_DOTS      = re.compile(r'[\.…]{3,}(?:(?:[ \t]{1,3}|\n\s*|/|\s*/\s*)[\.…]{3,})*')  # Match normal dots OR unicode ellipsis, including spaces
RE_CHECKBOX  = re.compile(r'\|_\|')
RE_PARENS    = re.compile(r'\(([^()]{2,200})\)')
RE_ONLY_PAREN = re.compile(r'^\s*\(([^()]{2,200})\)\s*$')

CHECKBOX_GROUP_GAP = 8  # max seq-number gap for grouping consecutive checkboxes


# ── Utilities ────────────────────────────────────────────────────────────────
def _sha(s: str) -> str:
    return hashlib.sha256(s.encode('utf-8')).hexdigest()[:16]


def _fold_ro(s: str) -> str:
    for src, dst in [('ă','a'),('â','a'),('î','i'),('ș','s'),('ş','s'),('ț','t'),('ţ','t')]:
        s = s.replace(src, dst)
    return s


def _norm(s: str) -> str:
    return re.sub(r'\s+', ' ', (s or '').strip())


def _is_hint_para(text: str) -> bool:
    """True if a paragraph is purely a label hint (no placeholders, short, has letters)."""
    t = _norm(text)
    if not t:
        return False
    if RE_ONLY_PAREN.match(t):
        return True
    if (len(t) <= 150
            and not RE_UND.search(t)
            and not RE_DOTS.search(t)
            and not RE_CHECKBOX.search(t)
            and re.search(r'[A-Za-zĂÂÎȘȚăâîșț]', t)):
        return True
    return False


def _detect_placeholders(text: str) -> List[Tuple[str, str, int, int]]:
    """
    Find all non-overlapping placeholders in text.
    Returns list of (type, raw_span, start, end).
    Priority: DATE > UND > DOTS > CHECKBOX.
    Must mask found items to prevent lower-priority regexes from spanning across boundaries.
    """
    chosen: List[Tuple[str, str, int, int]] = []
    
    # We will mutable a copy of the text, masking out found regions with spaces
    # so that subsequent regexes don't accidentally match half-inside and half-outside
    # a higher priority placeholder (like a Date followed by a newline).
    mask_text = bytearray(text, 'utf-8')
    def _mask_region(s: int, e: int):
        # bytearray indices might differ from char indices if there's unicode, 
        # so it's safer to just mask the string char-by-char or be careful.
        pass
        
    masked_str = text
    
    for ftype, pattern in [('DATE', RE_DATE), ('TEXT', RE_UND), ('TEXT', RE_DOTS), ('CHECKBOX', RE_CHECKBOX)]:
        for m in pattern.finditer(masked_str):
            s, e = m.start(), m.end()
            # If it's just spaces, ignore
            if masked_str[s:e].strip() == '':
                continue
            chosen.append((ftype, text[s:e], s, e))
            # Mask out the matched region with spaces to prevent overlap
            masked_str = masked_str[:s] + ' ' * (e - s) + masked_str[e:]

    chosen.sort(key=lambda x: x[2])
    return chosen


def _extract_label(
    full_text: str,
    prefix: str,
    suffix: str,
    ftype: str,
    next_text: Optional[str],
) -> Tuple[str, str]:
    """
    Extract the best label for a placeholder.
    Returns (label, source_description).
    """
    # 1. (hint) in parentheses right after the placeholder
    m = RE_PARENS.search(suffix)
    if m:
        lbl = _norm(m.group(1)).strip(' .;,:')
        if len(lbl) >= 2:
            return lbl, 'parens_suffix'

    # 2. (hint) anywhere in the full paragraph
    for m in RE_PARENS.finditer(full_text):
        lbl = _norm(m.group(1)).strip(' .;,:')
        if len(lbl) >= 2 and re.search(r'[A-Za-zĂÂÎȘȚăâîșț]', lbl):
            return lbl, 'parens_inline'

    # 3. CHECKBOX: option text is the text that follows |_| on the same line.
    if ftype == 'CHECKBOX':
        suf = _norm(suffix).strip()
        first_line = suf.splitlines()[0].strip(' .;,:') if suf else ''
        if len(first_line) >= 2:
            return first_line, 'checkbox_option'

    # 4. Next paragraph as a label hint (non-checkbox fields only)
    if next_text:
        # PRIORITIZE: If the next paragraph is ONLY a bracketed hint (e.g. "(denumirea serviciilor)"),
        # we take it as the label even if the prefix is long.
        m_only = re.match(r'^\s*\(\s*([^\)]+)\s*\)\s*$', next_text.strip())
        if m_only:
            t = _norm(m_only.group(1)).strip(' .;,:')
            if len(t) >= 2:
                return t, 'next_para_hint'

        # Fallback: only use next_para if prefix is short
        if _norm(prefix or '') == '' or len(_norm(prefix)) < 25:
            m3 = RE_PARENS.search(_norm(next_text))
            if m3:
                t = _norm(m3.group(1)).strip(' .;,:')
            else:
                t = _norm(next_text).strip(' .;,:()_.')
            if len(t) >= 2:
                return t, 'next_para'

    # 5. Last meaningful chunk of the text before the placeholder
    p = RE_UND.sub(' ', prefix or '')
    p = RE_DOTS.sub(' ', p)
    p = RE_CHECKBOX.sub(' ', p)
    for sep in ['\n', '\r', ';', '.']:
        parts = p.split(sep)
        p = parts[-1]
    for sep in [',', '–', '—']:
        parts = p.split(sep)
        p = parts[-1]
    p = _norm(p).strip(' .;,:_')
    words = p.split()
    if len(words) > 8:
        p = ' '.join(words[-6:])
    if len(p) >= 2 and re.search(r'[A-Za-zĂÂÎȘȚăâîșț]', p):
        return p, 'prefix'

    # 6. Fallback: first meaningful phrase from suffix (suffix_start)
    if suffix:
        s = RE_UND.sub(' ', RE_DOTS.sub(' ', suffix))
        s = re.sub(r'^[\s,;:()+\-–—]+', '', _norm(s))
        s = s.split(',')[0].split(';')[0].split('\n')[0]
        s = _norm(s).strip(' .;,:_()')
        words_s = s.split()
        if len(words_s) > 6:
            s = ' '.join(words_s[:5])
        if len(s) >= 4 and re.search(r'[A-Za-zĂÂÎȘȚăâîșț]', s):
            return s, 'suffix_start'

    return '', 'none'


# ── Data structures ───────────────────────────────────────────────────────────
@dataclass
class Field:
    field_id: str
    field_type: str       # TEXT | DATE | CHECKBOX
    location: str         # deterministic path, e.g. "body/p:5"
    seq: int              # global sequence number (used for checkbox grouping)
    full_text: str        # full paragraph text
    start: int            # placeholder start offset in full_text
    end: int              # placeholder end offset in full_text
    raw_span: str         # the actual placeholder characters
    label: str
    label_source: str
    ctx_before: str       # trimmed text before placeholder (same paragraph)
    ctx_after: str        # trimmed text after placeholder (same paragraph)
    ctx_prev_para: str    # text from previous paragraph(s)
    ctx_next_para: str    # text from next paragraph(s)
    group_id: Optional[str] = None    # set for CHECKBOX fields that form a radio group
    option_label: str = ''            # for CHECKBOX: this specific option's label

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class TableMeta:
    """Metadata for a table that should be filled from a JSON array."""
    field_id: str
    table_index: int
    location: str
    label: str              # from paragraph preceding the table
    col_headers: List[str]  # column header texts (from header row)
    data_rows: List[int]    # row indices that are empty/fillable

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


# ── Document iteration ────────────────────────────────────────────────────────
def iter_block_items(parent) -> Iterator:
    """Yield Paragraphs and Tables in document order for a Document or Cell."""
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(type(parent))
    for child in parent_elm.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            yield Paragraph(child, parent)
        elif tag == 'tbl':
            yield Table(child, parent)


def _compute_fingerprint(doc: DocxDocument, path: str) -> str:
    parts = [path]
    for p in list(doc.paragraphs)[:60]:
        parts.append(p.text or '')
    return _sha('\n'.join(parts))


def _parse_fields_from_text(
    text: str,
    location: str,
    seq: int,
    template_fp: str,
    next_text: Optional[str],
    col_header: str = '',
    prev_text: Optional[str] = None,
) -> List[Field]:
    """Extract Field objects from a single paragraph's text."""
    fields = []
    prev_ctx = _norm(prev_text or '')[-300:] if prev_text else ''
    next_ctx = _norm(next_text or '')[:300] if next_text else ''
    for ftype, span, s, e in _detect_placeholders(text):
        prefix, suffix = text[:s], text[e:]
        label, lsrc = _extract_label(text, prefix, suffix, ftype, next_text)

        if not label and col_header:
            label, lsrc = col_header, 'col_header'

        # field_id: NO label in hash — stable for LLM enrichment
        fid = _sha(f"{template_fp}|{location}|{ftype}|{s}|{e}|{span}")
        fields.append(Field(
            field_id=fid,
            field_type=ftype,
            location=location,
            seq=seq,
            full_text=text,
            start=s,
            end=e,
            raw_span=span,
            label=label,
            label_source=lsrc,
            ctx_before=_norm(prefix)[-200:],
            ctx_after=_norm(suffix)[:200],
            ctx_prev_para=prev_ctx,
            ctx_next_para=next_ctx,
        ))
    return fields


def _is_empty_cell(text: str) -> bool:
    t = text.strip().replace('\xa0', '').strip()
    return t == '' or all(c in ('_', '.', ' ', '\t', '\n') for c in t)


def _is_title_like(text: str) -> bool:
    """True if a paragraph looks like a section title (short, no placeholders)."""
    t = _norm(text)
    if not t or len(t) > 250:
        return False
    if RE_UND.search(t) or RE_DOTS.search(t) or RE_CHECKBOX.search(t):
        return False
    return bool(re.search(r'[A-Za-zĂÂÎȘȚăâîșț]', t))


def _parse_table_meta(
    table: Table,
    table_idx: int,
    recent_paras: List[str],
    template_fp: str,
) -> Optional[TableMeta]:
    """Build a TableMeta if the table has empty data rows."""
    rows = table.rows
    if len(rows) < 2:
        return None

    first_empty_row: Optional[int] = None
    for ri, row in enumerate(rows):
        if all(_is_empty_cell(cell.text) for cell in row.cells):
            first_empty_row = ri
            break

    if first_empty_row is None:
        return None

    header_row_idx: Optional[int] = None
    for ri in range(first_empty_row - 1, -1, -1):
        if any(not _is_empty_cell(rows[ri].cells[ci].text)
               for ci in range(len(rows[ri].cells))):
            header_row_idx = ri
            break

    if header_row_idx is None:
        return None

    col_headers = [_norm(rows[header_row_idx].cells[ci].text)
                   for ci in range(len(rows[header_row_idx].cells))]

    data_rows = []
    for ri, row in enumerate(rows):
        if ri <= header_row_idx:
            continue
        cells = [_norm(cell.text) for cell in row.cells]
        empty_count = sum(1 for t in cells if _is_empty_cell(t))
        if empty_count >= max(1, int(len(cells) * 0.8)):
            data_rows.append(ri)

    if not data_rows:
        return None

    if len(col_headers) < 3:
        return None

    label = f'table_{table_idx}'
    for p in reversed(recent_paras):
        if _is_title_like(p):
            label = _norm(p)[-200:]
            break

    fid = _sha(f"{template_fp}|table|{table_idx}|{label}")
    return TableMeta(
        field_id=fid,
        table_index=table_idx,
        location=f'body/t:{table_idx}',
        label=label,
        col_headers=col_headers,
        data_rows=data_rows,
    )


def _group_checkboxes(fields: List[Field], template_fp: str) -> None:
    """Assign group_id to consecutive CHECKBOX fields."""
    cb = [f for f in fields if f.field_type == 'CHECKBOX']
    if not cb:
        return

    groups: List[List[Field]] = []
    current = [cb[0]]
    for prev, curr in zip(cb, cb[1:]):
        if curr.seq - prev.seq <= CHECKBOX_GROUP_GAP:
            current.append(curr)
        else:
            groups.append(current)
            current = [curr]
    groups.append(current)

    for group in groups:
        if len(group) < 2:
            continue
        gid = _sha(f"{template_fp}|cbgroup|{group[0].location}|{group[0].start}")
        for f in group:
            f.group_id = gid
            f.option_label = f.label


# ── Public API ────────────────────────────────────────────────────────────────
def parse_document(docx_path: str) -> Dict[str, Any]:
    """
    Parse a DOCX form and return all fillable fields.

    Returns:
        {
            "template_fingerprint": str,
            "fields": [Field.to_dict(), ...],
            "tables": [TableMeta.to_dict(), ...]
        }
    """
    doc = Document(docx_path)
    template_fp = _compute_fingerprint(doc, docx_path)

    fields: List[Field] = []
    tables: List[TableMeta] = []

    body_blocks = list(iter_block_items(doc))
    global_seq = 0
    table_idx_counter = 0
    para_idx_in_body = 0
    recent_body_paras: List[str] = []

    for bi, block in enumerate(body_blocks):
        if isinstance(block, Paragraph):
            text = block.text or ''

            prev_text: Optional[str] = None
            for k in range(bi - 1, -1, -1):
                if isinstance(body_blocks[k], Paragraph):
                    prev_text = body_blocks[k].text or ''
                    break

            next_text: Optional[str] = None
            for k in range(bi + 1, len(body_blocks)):
                if isinstance(body_blocks[k], Paragraph):
                    next_text = body_blocks[k].text or ''
                    break

            location = f'body/p:{para_idx_in_body}'
            new_fields = _parse_fields_from_text(
                text, location, global_seq, template_fp, next_text, prev_text=prev_text
            )
            fields.extend(new_fields)
            global_seq += 1
            if text.strip():
                recent_body_paras.append(text)
                if len(recent_body_paras) > 10:
                    recent_body_paras.pop(0)
            para_idx_in_body += 1

        elif isinstance(block, Table):
            tm = _parse_table_meta(block, table_idx_counter, recent_body_paras, template_fp)
            if tm:
                tables.append(tm)

            for ri, row in enumerate(block.rows):
                col_hdrs = [_norm(block.rows[0].cells[ci].text)
                            for ci in range(len(row.cells))]
                for ci, cell in enumerate(row.cells):
                    cell_blocks = list(iter_block_items(cell))
                    for pi, cp in enumerate(cell_blocks):
                        if not isinstance(cp, Paragraph):
                            continue
                        ctext = cp.text or ''
                        prev_cell_text: Optional[str] = None
                        if pi > 0 and isinstance(cell_blocks[pi - 1], Paragraph):
                            prev_cell_text = cell_blocks[pi - 1].text or ''
                        next_cell_text: Optional[str] = None
                        for k in range(pi + 1, len(cell_blocks)):
                            if isinstance(cell_blocks[k], Paragraph):
                                next_cell_text = cell_blocks[k].text or ''
                                break
                        location = f'body/t:{table_idx_counter}/r:{ri}/c:{ci}/p:{pi}'
                        col_hdr = col_hdrs[ci] if ci < len(col_hdrs) else ''
                        new_fields = _parse_fields_from_text(
                            ctext, location, global_seq, template_fp, next_cell_text, col_hdr,
                            prev_text=prev_cell_text,
                        )
                        fields.extend(new_fields)
                        global_seq += 1

            table_idx_counter += 1

    for si, section in enumerate(doc.sections):
        header_paras = section.header.paragraphs
        for pi, p in enumerate(header_paras):
            text = p.text or ''
            prev_h = header_paras[pi - 1].text if pi > 0 else None
            location = f'sec:{si}/header/p:{pi}'
            new_fields = _parse_fields_from_text(
                text, location, global_seq, template_fp, None, prev_text=prev_h
            )
            fields.extend(new_fields)
            global_seq += 1

        footer_paras = section.footer.paragraphs
        for pi, p in enumerate(footer_paras):
            text = p.text or ''
            prev_f = footer_paras[pi - 1].text if pi > 0 else None
            location = f'sec:{si}/footer/p:{pi}'
            new_fields = _parse_fields_from_text(
                text, location, global_seq, template_fp, None, prev_text=prev_f
            )
            fields.extend(new_fields)
            global_seq += 1

    _group_checkboxes(fields, template_fp)

    return {
        'template_fingerprint': template_fp,
        'fields': [f.to_dict() for f in fields],
        'tables': [t.to_dict() for t in tables],
    }


if __name__ == '__main__':
    import argparse
    import json

    ap = argparse.ArgumentParser(description='Parse DOCX form and output fields + tables as JSON')
    ap.add_argument('--docx', required=True, help='Path to input DOCX')
    ap.add_argument('--out', default='parser_result.json', help='Output JSON path')
    args = ap.parse_args()

    result = parse_document(args.docx)
    with open(args.out, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f'Fields: {len(result["fields"])}  Tables: {len(result["tables"])}')
    print(f'Output: {args.out}')
